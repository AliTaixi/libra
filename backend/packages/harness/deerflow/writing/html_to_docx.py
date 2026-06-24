"""HTML → .docx 直转器。

将 wangEditor 输出的完整 HTML（含行内样式）直接转换为 Word 文档，
保留字体、颜色、字号、对齐、表格、列表、图片等全部格式。

依赖：bs4, lxml, python-docx（均已预装）
"""

from __future__ import annotations

import logging
import re
from io import BytesIO

from bs4 import BeautifulSoup, NavigableString, Tag

logger = logging.getLogger("deerflow.writing.html_to_docx")
from docx import Document as _DocumentFn
from docx.document import Document as _DocumentClass
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm


# ── 样式解析 ──────────────────────────────────────────────────────────────

def _parse_css_style(style: str | None) -> dict[str, str]:
    if not style:
        return {}
    result = {}
    for part in style.split(";"):
        part = part.strip()
        if ":" in part:
            key, val = part.split(":", 1)
            result[key.strip().lower()] = val.strip()
    return result


def _parse_color(color: str) -> RGBColor | None:
    color = color.strip()
    if color.startswith("#"):
        c = color.lstrip("#")
        if len(c) == 3:
            c = "".join(x * 2 for x in c)
        try:
            return RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
        except (ValueError, IndexError):
            return None
    if color.startswith("rgb"):
        parts = re.findall(r"\d+", color)
        if len(parts) >= 3:
            return RGBColor(int(parts[0]), int(parts[1]), int(parts[2]))
    named = {
        "red": RGBColor(0xFF, 0, 0), "black": RGBColor(0, 0, 0),
        "white": RGBColor(0xFF, 0xFF, 0xFF), "blue": RGBColor(0, 0, 0xFF),
        "green": RGBColor(0, 0x80, 0), "gray": RGBColor(0x80, 0x80, 0x80),
        "yellow": RGBColor(0xFF, 0xFF, 0),
    }
    return named.get(color.lower())


def _parse_font_size(size_str: str) -> float | None:
    size_str = size_str.strip().lower()
    m = re.match(r"([\d.]+)\s*(px|pt|em)?", size_str)
    if not m:
        return None
    val = float(m.group(1))
    unit = m.group(2) or "pt"
    return val * {"px": 0.75, "em": 12}.get(unit, 1)


def _resolve_font_family(font_str: str | None) -> tuple[str, str]:
    if not font_str:
        return "宋体", "Times New Roman"
    fonts = [f.strip().strip("'\"") for f in font_str.split(",")]
    for f in fonts:
        fl = f.lower()
        if fl in ("宋体", "simsun"):
            return "宋体", "Times New Roman"
        if fl in ("微软雅黑", "microsoft yahei", "yahei"):
            return "微软雅黑", "Times New Roman"
        if fl in ("黑体", "simhei"):
            return "黑体", "Times New Roman"
        if fl in ("楷体", "kaiti"):
            return "楷体", "Times New Roman"
        if fl in ("times new roman",):
            return "宋体", "Times New Roman"
        if fl in ("arial",):
            return "宋体", "Arial"
        if fl in ("calibri",):
            return "宋体", "Calibri"
    return "宋体", "Times New Roman"


# ── python-docx 辅助 ──────────────────────────────────────────────────────

def _set_run_font(run, cn_font="宋体", en_font="Times New Roman",
                  size: float | None = None,
                  bold: bool | None = None, italic: bool | None = None,
                  color: RGBColor | None = None,
                  underline: bool | None = None, strike: bool | None = None):
    run.font.name = en_font
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color: run.font.color.rgb = color
    if underline: run.font.underline = underline
    if strike: run.font.strike = strike
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def _set_align(p, align: str | None):
    if not align:
        return
    m = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = m.get(align.lower())


def _add_shading(element, color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color.lstrip("#"))
    # 找到正确的插入位置
    for tag in ("w:tcPr", "w:pPr"):
        c = element.find(qn(tag))
        if c is not None:
            c.append(shd)
            return
    # 没有就新建
    parent_tag = element.tag
    new_tag = "w:pPr" if "p" in parent_tag else "w:tcPr"
    container = OxmlElement(new_tag)
    container.append(shd)
    element.insert(0, container)


HEADING_SIZES = {1: 22, 2: 16, 3: 14, 4: 12}


# ── 累积行内样式 ──────────────────────────────────────────────────────────

class Style:
    """在嵌套标签间传递的累积样式。"""
    def __init__(self, **kw):
        self.bold = kw.get("bold")
        self.italic = kw.get("italic")
        self.underline = kw.get("underline")
        self.strike = kw.get("strike")
        self.cn = kw.get("cn", "宋体")
        self.en = kw.get("en", "Times New Roman")
        self.size: float | None = kw.get("size")
        self.color: RGBColor | None = kw.get("color")
        self.bg: str | None = kw.get("bg")
        self.align: str | None = kw.get("align")

    def with_css(self, css: dict[str, str]) -> "Style":
        s = Style(bold=self.bold, italic=self.italic, underline=self.underline,
                  strike=self.strike, cn=self.cn, en=self.en, size=self.size,
                  color=self.color, bg=self.bg, align=self.align)
        if "font-family" in css:
            s.cn, s.en = _resolve_font_family(css["font-family"])
        if "font-size" in css:
            s.size = _parse_font_size(css["font-size"])
        if "color" in css:
            s.color = _parse_color(css["color"])
        if "background-color" in css:
            s.bg = css["background-color"]
        if "text-align" in css:
            s.align = css["text-align"]
        return s

    def with_tag(self, tag: str) -> "Style":
        s = Style(bold=self.bold, italic=self.italic, underline=self.underline,
                  strike=self.strike, cn=self.cn, en=self.en, size=self.size,
                  color=self.color, bg=self.bg, align=self.align)
        if tag in ("strong", "b"): s.bold = True
        elif tag in ("em", "i"): s.italic = True
        elif tag == "u": s.underline = True
        elif tag in ("s", "del"): s.strike = True
        elif tag == "code": s.en = "Courier New"; s.size = 9
        elif tag in ("th",): s.bold = True
        return s

    def apply(self, run):
        _set_run_font(run, cn_font=self.cn, en_font=self.en, size=self.size,
                      bold=self.bold, italic=self.italic, color=self.color,
                      underline=self.underline, strike=self.strike)


# ── 核心转换 ──────────────────────────────────────────────────────────────


def html_to_docx(html: str) -> bytes:
    """HTML → .docx 字节流。"""
    doc = _DocumentFn()
    _set_run_font(doc.styles["Normal"], size=12)
    for lv in range(1, 5):
        hs = f"Heading {lv}"
        if hs in doc.styles:
            _set_run_font(doc.styles[hs], size=HEADING_SIZES[lv], bold=True, color=RGBColor(0, 0, 0))

    soup = BeautifulSoup(html, "lxml")
    body = soup.find("body") or soup
    _render(doc, body, Style())

    import tempfile
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    with open(path, "rb") as f:
        data = f.read()
    import os
    os.unlink(path)
    return data


def _render(parent, node, style: Style):
    """递归渲染节点。"""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                if isinstance(parent, _DocumentClass):
                    p = parent.add_paragraph()
                    style.apply(p.add_run(text))
                else:
                    style.apply(parent.add_run(text))
            continue
        if not isinstance(child, Tag):
            continue
        tag = child.name
        s = style.with_css(_parse_css_style(child.get("style", "")))

        if tag in ("h1", "h2", "h3", "h4"):
            lv = int(tag[1])
            p = parent.add_heading(level=lv)
            _set_align(p, s.align)
            _render(p, child, s)
            for r in p.runs:
                r.font.bold = True

        elif tag == "p":
            p = parent.add_paragraph()
            _set_align(p, s.align)
            if s.bg: _add_shading(p._element, s.bg)
            p.paragraph_format.first_line_indent = Cm(0.74)
            _render(p, child, s)

        elif tag in ("ul", "ol"):
            for li in child.find_all("li", recursive=False):
                p = parent.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
                _render(p, li, s)

        elif tag == "table":
            _render_table(parent, child, s)

        elif tag == "blockquote":
            p = parent.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            pPr = p._element.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "999999")
            bdr.append(left)
            pPr.append(bdr)
            _render(p, child, s)

        elif tag == "pre":
            p = parent.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            _add_shading(p._element, "f0f0f0")
            code_s = Style(cn="宋体", en="Courier New", size=9)
            code_s.apply(p.add_run(child.get_text("\n")))

        elif tag in ("figure", "div", "thead", "tbody", "tfoot"):
            _render(parent, child, s)

        elif tag == "hr":
            p = parent.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:space"), "1")
            b.set(qn("w:color"), "999999")
            bdr.append(b)
            pPr.append(bdr)

        elif tag == "img":
            src = child.get("src", "")
            import base64
            import urllib.request
            import urllib.error
            from pathlib import Path

            # ── 获取图片来源 ──
            picture_source: str | BytesIO | None = None

            try:
                if src.startswith("file://"):
                    img_bytes = Path(src[7:]).read_bytes()
                    picture_source = BytesIO(img_bytes)
                elif src.startswith("data:"):
                    _, encoded = src.split(",", 1)
                    img_bytes = base64.b64decode(encoded)
                    picture_source = BytesIO(img_bytes)
                elif Path(src).exists():
                    picture_source = src
                else:
                    img_bytes = urllib.request.urlopen(src, timeout=15).read()
                    picture_source = BytesIO(img_bytes)

                if picture_source:
                    # parent 可能是 Document 或 Paragraph（img 嵌套在 p 内时）
                    if isinstance(parent, _DocumentClass):
                        p = parent.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(picture_source, width=Inches(5))
                    else:
                        # parent 是 Paragraph → 在当前段落添加图片
                        parent.add_run().add_picture(picture_source, width=Inches(5))
            except Exception as exc:
                logger.warning("图片加载失败: src=%s error=%s", src[:120], exc)

        elif tag == "figcaption":
            p = parent.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _render(p, child, s.with_css({"font-size": "9pt", "color": "#666666"}))

        elif tag == "br":
            if isinstance(parent, _DocumentClass):
                parent.add_paragraph()
            else:
                parent.add_run("\n")

        elif tag == "a":
            text = child.get_text()
            if text:
                run = (parent.add_run(text) if not isinstance(parent, _DocumentClass)
                       else parent.add_paragraph().add_run(text))
                s.with_css({"color": "#0563C1", "font-family": "Calibri"}).apply(run)
                run.font.underline = True

        elif tag in ("span", "code", "strong", "b", "em", "i", "u", "s", "del", "sub", "sup"):
            _render(parent, child, s.with_tag(tag))

        else:
            _render(parent, child, s)


def _render_table(parent, table_tag: Tag, style: Style):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    max_cols = max((len(r.find_all(["th", "td"])) for r in rows), default=0)
    if max_cols == 0:
        return
    table = parent.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for ri, row_tag in enumerate(rows):
        for ci, cell_tag in enumerate(row_tag.find_all(["th", "td"])):
            if ci >= max_cols:
                break
            cell = table.rows[ri].cells[ci]
            is_h = cell_tag.name == "th"
            cs = style.with_css(_parse_css_style(cell_tag.get("style", "")))
            if is_h:
                _add_shading(cell._element, "E0E0E0")
                cs = cs.with_tag("th")
            if cs.bg:
                _add_shading(cell._element, cs.bg)
            cell.text = ""
            _set_align(cell.paragraphs[0], cs.align)
            _render(cell.paragraphs[0], cell_tag, cs)
