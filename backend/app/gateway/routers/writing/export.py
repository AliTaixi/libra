"""文档导出 + 模板解析 API — 薄路由层，逻辑在 deerflow.writing.service。"""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import subprocess
import tempfile
import uuid
from pathlib import Path
from urllib.parse import quote

from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import Response

from .models import (
    ExportDocxRequest, ExportDocxResponse,
    ExportMdDocxRequest,
    ParseTemplateRequest, ParseTemplateResponse, ChapterInfo,
)
from pydantic import BaseModel, Field

from deerflow.writing.service import (
    WritingError,
    export_md_to_docx, export_html_to_docx,
    _user_data_dir, _ensure_output_dir,
)

logger = logging.getLogger("app.gateway.routers.writing")

router = APIRouter()


def _ok(e: WritingError) -> HTTPException:
    return HTTPException(status_code=e.status_code, detail=e.message)


@router.post("/export-md-docx")
async def export_md_docx(req: ExportMdDocxRequest):
    """Markdown → Word 并返回文件流。"""
    try:
        docx_bytes = await export_md_to_docx(req.markdown, req.output_name)
        from fastapi.responses import Response
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(req.output_name)}"},
        )
    except WritingError as e:
        raise _ok(e)


class ExportHtmlDocxRequest(BaseModel):
    """HTML → docx 导出请求。"""
    html: str = Field(..., description="完整文档的 HTML 内容")
    output_name: str = Field(default="output.docx", description="输出文件名")


def _resolve_img_src(html: str, cookies: dict[str, str]) -> str:
    """将 HTML 中指向本服务的图片 URL 替换为本地临时文件路径。

    用请求的 cookie 下载图片 → 存临时文件 → 替换 src 为文件路径，
    这样 html_to_docx 的 ``add_picture(路径)`` 可直接读取。
    """
    import urllib.request

    def _download(src: str) -> str | None:
        """下载图片到临时文件，返回文件路径。"""
        try:
            # 相对路径 → 补全为后端地址（前端通过 Next.js rewrite 转发，
            # 后端内部直接调后端自己）
            if src.startswith("/"):
                gw = os.environ.get("DEER_FLOW_INTERNAL_GATEWAY_BASE_URL", "http://127.0.0.1:8002")
                src = f"{gw}{src}"
            req = urllib.request.Request(src)
            # 注入 cookie 使后端能验证身份
            cookie_str = "; ".join(f"{k}={v}" for k, v in cookies.items())
            if cookie_str:
                req.add_header("Cookie", cookie_str)
            with urllib.request.urlopen(req, timeout=15) as resp:
                img_data = resp.read()
            # 根据 Content-Type 推断后缀
            ct = resp.headers.get("Content-Type", "image/png")
            ext = {"image/png": ".png", "image/jpeg": ".jpg", "image/gif": ".gif",
                   "image/webp": ".webp", "image/bmp": ".bmp"}.get(ct, ".png")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
            tmp.write(img_data)
            tmp.close()
            logger.info("图片下载成功: %s → %s  (%d bytes)", src[:80], tmp.name, len(img_data))
            return tmp.name
        except Exception as e:
            logger.warning("图片下载失败: %s  error=%s", src[:80], e)
            return None

    def _replacer(m: re.Match) -> str:
        src = m.group(2)
        local_path = _download(src)
        if local_path:
            return m.group(0).replace(f'src={m.group(1)}{src}{m.group(1)}', f'src="{local_path}"')
        return m.group(0)

    return re.sub(r'<img\s+[^>]*src=(["\'])([^"\']+)\1', _replacer, html)


@router.post("/export-html-docx")
async def export_html_docx(req: ExportHtmlDocxRequest, request: Request):
    """HTML → docx 直转（保留字体/颜色/对齐/表格等全部格式）。"""
    try:
        # 提取请求中的 cookie，用于下载图片时验证身份
        cookies = dict(request.cookies)
        resolved_html = _resolve_img_src(req.html, cookies)

        docx_bytes = await asyncio.to_thread(export_html_to_docx, resolved_html)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(req.output_name)}"},
        )
    except WritingError as e:
        raise _ok(e)
    except Exception as e:
        logger.exception("HTML→docx 导出失败")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export-docx", response_model=ExportDocxResponse)
async def export_docx(req: ExportDocxRequest):
    """在原始模板上注入内容（逻辑在原 writing.py 导出部分，逐步迁移中）。"""
    # 当前实现保留在原处，暂时直接转发到 md-to-docx 降级路径
    # TODO: 将 inject 脚本逻辑也下沉到 harness
    from deerflow.writing.service import SKILLS_DIR
    import json, subprocess

    template_path = Path(req.template_path)
    _ensure_output_dir()

    if not template_path.exists():
        if req.markdown:
            docx_bytes = await export_md_to_docx(req.markdown, req.output_name)
            disk_name = f"{uuid.uuid4().hex}.docx"
            final_path = _user_data_dir() / disk_name
            final_path.parent.mkdir(parents=True, exist_ok=True)
            final_path.write_bytes(docx_bytes)
            return ExportDocxResponse(file_path=str(final_path), success=True)
        raise HTTPException(status_code=404, detail=f"模板不存在: {req.template_path}")

    out_path = _user_data_dir() / req.output_name
    script = _user_data_dir() / f"inject_{uuid.uuid4().hex}.py"
    chapters_json = json.dumps(
        [{"title": ch.title, "content": ch.content} for ch in req.chapters],
        ensure_ascii=False,
    )
    r_template = str(template_path).replace("\\", "/")
    r_out = str(out_path).replace("\\", "/")
    script.write_text(f"""#!/usr/bin/env python3
import sys, json, re
from docx import Document
template_path = r'{r_template}'
out_path = r'{r_out}'
chapters_data = {chapters_json}
doc = Document(template_path)
paras = [(p.style.name if p.style else "", p.text.strip(), p) for p in doc.paragraphs]
h1 = [i for i,(s,_,_) in enumerate(paras) if "heading" in s.lower() and re.search(r"1$",s)]
if not h1:
    doc.save(out_path); print(f"done: {{out_path}}"); sys.exit(0)
for ch in chapters_data:
    for idx in h1:
        if paras[idx][1]==ch["title"]:
            end = next((i for i in h1 if i>idx), len(paras))
            bodies = [bi for bi in range(idx+1,end) if paras[bi][1] and "heading" not in paras[bi][0].lower() and (len(paras[bi][1])>30 or any(k in paras[bi][1] for k in ["本条","本文档"]))]
            if bodies and ch["content"]:
                import html as hm
                txt = hm.unescape(re.sub(r"<[^>]+>","\n",ch["content"]))
                lines = [l.strip() for l in txt.split("\n") if l.strip()]
                for li, pi in enumerate(bodies):
                    if li < len(lines): paras[pi][2].text = lines[li]
            break
doc.save(out_path)
print(f"done: {{out_path}}")
""")

    try:
        result = subprocess.run(["python3", str(script)], capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            logger.warning(f"注入失败: {result.stderr}")
            if req.markdown:
                docx_bytes = await export_md_to_docx(req.markdown, req.output_name)
                final_path = _user_data_dir() / f"{uuid.uuid4().hex}.docx"
                final_path.write_bytes(docx_bytes)
                return ExportDocxResponse(file_path=str(final_path), success=True)
            raise HTTPException(status_code=500, detail=f"文档生成失败: {result.stderr}")
        if not out_path.exists():
            raise HTTPException(status_code=500, detail="输出文件不存在")
        return ExportDocxResponse(file_path=str(out_path), success=True)
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=504, detail="导出超时")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("导出失败")
        if req.markdown:
            docx_bytes = await export_md_to_docx(req.markdown, req.output_name)
            final_path = _user_data_dir() / f"{uuid.uuid4().hex}.docx"
            final_path.write_bytes(docx_bytes)
            return ExportDocxResponse(file_path=str(final_path), success=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/parse-template", response_model=ParseTemplateResponse)
async def parse_template(req: ParseTemplateRequest):
    """解析 Word 模板。"""
    template_path = Path(req.template_path)
    if not template_path.exists():
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {req.template_path}")
    try:
        import importlib.util as _imp_util
        _spec = _imp_util.spec_from_file_location("docx_to_md", "/app/skills/public/docx-to-md/scripts/docx_to_md.py")
        _mod = _imp_util.module_from_spec(_spec)
        _spec.loader.exec_module(_mod)
        chapters_raw, _ = _mod.extract(template_path)
        chapters_list = [ChapterInfo(**ch) for ch in (chapters_raw or [])]
        if not chapters_list:
            chapters_list = [ChapterInfo(id="1", title="概述")]
        return ParseTemplateResponse(chapters=chapters_list, variables=[], success=True)
    except Exception as e:
        logger.exception("模板解析失败")
        return ParseTemplateResponse(
            chapters=[ChapterInfo(id="1", title="概述", sub_titles=[])],
            variables=[], success=True,
        )
