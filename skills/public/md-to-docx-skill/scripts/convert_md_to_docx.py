#!/usr/bin/env python3
"""Markdown → Word (.docx) 转换器"""
import argparse, os, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_font(run, cn="宋体", en="Times New Roman", size=12, bold=False, italic=False, color=None):
    """Set font on a run. All text forced to black by default."""
    run.font.size = Pt(size)
    run.font.name = en
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None: rPr = r.makeelement(qn("w:rPr"), {}); r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None: rFonts = rPr.makeelement(qn("w:rFonts"), {}); rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)


def set_paragraph_indent(paragraph, first_line_cm=0.74):
    """Set first-line indent (默认 0.74cm ≈ 2 个中文字符)."""
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(first_line_cm)

def add_code_run(p, text, size=9):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Courier New"
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "f0f0f0")
    shd.set(qn("w:val"), "clear")
    run._element.get_or_add_rPr().append(shd)

def _set_code_block_style(p):
    """给代码块段落加左边框、缩进和间距，与正文区分。"""
    from docx.shared import Pt as PtSize
    p.paragraph_format.space_before = PtSize(6)
    p.paragraph_format.space_after = PtSize(6)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    left = pBdr.makeelement(qn("w:left"), {
        qn("w:val"): "single",
        qn("w:sz"): "12",
        qn("w:space"): "6",
        qn("w:color"): "BBBBBB",
    })
    pBdr.append(left)
    pPr.append(pBdr)

def parse_inline(p, text):
    i = 0
    while i < len(text):
        if text[i:i+2] == "**":
            j = text.find("**", i+2)
            if j > i:
                run = p.add_run(text[i+2:j])
                set_font(run, bold=True)
                i = j + 2
                continue
        elif text[i:i+1] == "*":
            j = text.find("*", i+1)
            if j > i and (j+1 >= len(text) or text[j+1] != "*"):
                run = p.add_run(text[i+1:j])
                set_font(run, italic=True)
                i = j + 1
                continue
        elif text[i:i+1] == "`":
            j = text.find("`", i+1)
            if j > i:
                add_code_run(p, text[i+1:j])
                i = j + 1
                continue
        # Normal text
        j = i + 1
        while j < len(text) and text[j] not in "*`":
            # Check for ** or * at j
            if text[j] == "*":
                if j+1 < len(text) and text[j+1] == "*":
                    break
                break
            if text[j] == "`":
                break
            j += 1
        if j > i:
            run = p.add_run(text[i:j])
            set_font(run)
        i = j

def _sanitize_svg_fonts(svg_path: str) -> str:
    """Replace SVG font-family with environment-safe fallbacks so text renders.

    SVGs from tools like Mermaid use fonts (Trebuchet MS, Verdana, etc.)
    that may not be installed in the sandbox. We replace all ``font-family``
    declarations with ``sans-serif``, ``serif``, or ``monospace``.
    """
    with open(svg_path, "r", encoding="utf-8") as f:
        content = f.read()
    # Replace font-family in style tags and inline style attributes
    content = re.sub(
        r'font-family\s*:\s*[^;}\n]+',
        'font-family:sans-serif',
        content,
        flags=re.IGNORECASE,
    )
    # Also replace font-family="" attributes
    content = re.sub(
        r'font-family="[^"]*"',
        'font-family="sans-serif"',
        content,
        flags=re.IGNORECASE,
    )
    return content


def _add_image_to_docx(doc, paragraph, path: str) -> None:
    """Add an image (PNG/JPG/SVG) to a paragraph.

    For SVG files, tries in order:
    1. Direct embed via ``add_picture`` (works with python-docx >=1.1)
    2. Convert to PNG via cairosvg (with CJK font)
    3. Text placeholder
    """
    if path.lower().endswith(".svg"):
        import tempfile, subprocess
        tmp_png = os.path.join(tempfile.gettempdir(), os.urandom(8).hex() + ".png")

        # Write sanitised SVG (fonts → sans-serif) to a tmp file
        sanitised = _sanitize_svg_fonts(path)
        tmp_svg = os.path.join(tempfile.gettempdir(), os.urandom(8).hex() + ".svg")
        try:
            with open(tmp_svg, "w", encoding="utf-8") as f:
                f.write(sanitised)

            # Method 1: rsvg-convert (best quality, handles complex SVGs)
            try:
                subprocess.run(
                    ["rsvg-convert", "--dpi-x=300", "--dpi-y=300", "-o", tmp_png, tmp_svg],
                    check=True, capture_output=True, timeout=30,
                )
                if os.path.exists(tmp_png):
                    paragraph.add_run().add_picture(tmp_png, width=Inches(5))
                    return
            except Exception:
                pass

            # Method 2: cairosvg
            try:
                import cairosvg
                cairosvg.svg2png(url=tmp_svg, write_to=tmp_png, dpi=300)
                if os.path.exists(tmp_png):
                    paragraph.add_run().add_picture(tmp_png, width=Inches(5))
                    return
            except Exception:
                pass

            # Method 3: inkscape
            try:
                subprocess.run(
                    ["inkscape", "--without-gui", "--export-dpi=300",
                     "--export-filename=" + tmp_png, tmp_svg],
                    check=True, capture_output=True, timeout=30,
                )
                if os.path.exists(tmp_png):
                    paragraph.add_run().add_picture(tmp_png, width=Inches(5))
                    return
            except Exception:
                pass
        finally:
            for p in [tmp_svg, tmp_png]:
                try:
                    os.remove(p)
                except OSError:
                    pass

        # Final fallback: text placeholder
        run = paragraph.add_run(f"[SVG: {os.path.basename(path)}]")
        set_font(run, size=10, italic=True)
    else:
        # PNG/JPG/GIF etc — direct insert
        paragraph.add_run().add_picture(path, width=Inches(5))


def _apply_base_font(style, cn="宋体", en="Times New Roman", size=12):
    """Apply font settings to a document style (paragraph-level)."""
    style.font.name = en
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)


HEADING_SIZES = {1: 22, 2: 16, 3: 14, 4: 12, 5: 10, 6: 10}


def md_to_docx(md_text, output_path):
    doc = Document()
    # 设置默认字体（正文）
    _apply_base_font(doc.styles["Normal"], size=12)
    # 设置所有标题样式字体（宋体 + Times New Roman）
    for level in range(1, 7):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            _apply_base_font(doc.styles[style_name], size=HEADING_SIZES[level])

    lines = md_text.split("\n")
    i = 0
    code_block = False
    code_buf = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith("```"):
            if code_block:
                p = doc.add_paragraph()
                _set_code_block_style(p)
                for cl in code_buf:
                    add_code_run(p, cl + "\n")
                code_block = False
                code_buf = []
                i += 1
                continue
            else:
                # Read language if any
                code_block = True
                i += 1
                continue

        if code_block:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line):
            doc.add_paragraph("").paragraph_format.space_before = Pt(6)
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "6",
                qn("w:space"): "1",
                qn("w:color"): "999999",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        h_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                set_font(run, size=HEADING_SIZES[level], bold=True)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            ind = pPr.makeelement(qn("w:ind"), {
                qn("w:left"): "720",
                qn("w:hanging"): "0",
            })
            pPr.append(ind)
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            left = pBdr.makeelement(qn("w:left"), {
                qn("w:val"): "single",
                qn("w:sz"): "12",
                qn("w:space"): "8",
                qn("w:color"): "999999",
            })
            pBdr.append(left)
            pPr.append(pBdr)
            parse_inline(p, text)
            i += 1
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip(" |").split("|")]
            # 跳过分隔行（---|---）
            if all(re.match(r"^[-:\s]+$", c) for c in cells):
                i += 1
                continue
            if not in_table:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                for ci, c in enumerate(cells):
                    cell = table.rows[0].cells[ci]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(c)
                    set_font(run, size=10, bold=True)
                    # 表头灰色背景 + 黑色文字
                    shading = cell._element.get_or_add_tcPr().makeelement(qn("w:shd"), {
                        qn("w:val"): "clear",
                        qn("w:fill"): "E0E0E0",
                    })
                    cell._element.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                in_table = True
            else:
                row = table.add_row()
                for ci, c in enumerate(cells):
                    if ci < len(row.cells):
                        cell = row.cells[ci]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(c)
                        set_font(run, size=10)
            i += 1
            continue
        else:
            if in_table:
                # 检查下一行是否还是表格
                if i+1 < len(lines) and lines[i+1].strip().startswith("|"):
                    # Separator row, skip
                    pass
                in_table = False

        # Empty line
        if not line.strip():
            i += 1
            continue

        # List (unordered)
        ul_match = re.match(r"^[\s]*[-*+]\s+(.+)$", line)
        if ul_match:
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, ul_match.group(1))
            i += 1
            continue

        # List (ordered)
        ol_match = re.match(r"^[\s]*\d+\.\s+(.+)$", line)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, ol_match.group(1))
            i += 1
            continue

        # Image (支持文件名中含括号，如 "image (1).svg")
        img_match = re.match(r"!\[([^\]]*)\]\(((?:[^()]|\([^()]*\))*)\)", line)
        if img_match:
            caption = img_match.group(1)
            path = img_match.group(2)
            abs_path = os.path.abspath(path)
            if os.path.exists(abs_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_image_to_docx(doc, p, abs_path)
                    if caption:
                        cp = doc.add_paragraph()
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cp.add_run(f"图：{caption}")
                        set_font(run, size=9)
                except Exception as e:
                    # Fallback: insert caption text instead of image
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"[图片: {caption or os.path.basename(abs_path)}]")
                    set_font(run, size=10, italic=True)
            i += 1
            continue

        # Regular paragraph（首行缩进）
        p = doc.add_paragraph()
        set_paragraph_indent(p)
        parse_inline(p, line)
        i += 1

    # 如果代码块到文件末尾还没关闭，强制输出
    if code_block and code_buf:
        p = doc.add_paragraph()
        _set_code_block_style(p)
        for cl in code_buf:
            add_code_run(p, cl + "\n")

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    print(f"文档已生成: {output_path}")

if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--input", required=True)
    p.add_argument("--output", default="/mnt/user-data/outputs/输出.docx")
    args = p.parse_args()
    with open(args.input, "r", encoding="utf-8") as f:
        md_text = f.read()
    md_to_docx(md_text, args.output)
